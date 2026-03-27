from docx import Document
import os

class DOCXParser:
    def __init__(self, logger=None):
        self.logger = logger

    def extract_text(self, docx_path: str) -> str:
        """Extracts text from a DOCX, including paragraphs and tables."""
        if not os.path.exists(docx_path):
            if self.logger:
                self.logger.error(f"File not found: {docx_path}")
            return ""

        full_text = []
        try:
            doc = Document(docx_path)
            
            # 1. Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # 2. Extract tables
            for table in doc.tables:
                full_text.append("\n[Table Data]")
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    full_text.append(" | ".join(row_data))
                full_text.append("[End Table Data]\n")

            return "\n".join(full_text)
        except Exception as e:
            if self.logger:
                self.logger.error(f"Error parsing DOCX {docx_path}: {e}")
            return ""
